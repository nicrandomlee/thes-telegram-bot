import docx
from src.utils.gspread_utils import get_list_of_befriending_seniors_status, get_list_of_frail_seniors_status
from src.utils.utils import extract_frail_senior_name_from_frail_list
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, top=True, right=True, bottom=True, left=True):
    """
    Set cell borders
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # List of all borders
    borders = [
        ('top', top),
        ('right', right),
        ('bottom', bottom),
        ('left', left)
    ]

    for border, value in borders:
        if value:
            tag = f'w:{border}'
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto')

def add_table_with_grid(doc, rows, cols):
    table = doc.add_table(rows=rows, cols=cols)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell)
    return table

def generate_report_from_template(template_path):
    doc = docx.Document(template_path)

    doc.add_page_break()

    title = doc.add_paragraph("Befriending Seniors' House Visit Updates")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    befriending_seniors_list = get_list_of_befriending_seniors_status()[1:]
    befriending_seniors_list = list(map(lambda row: row[:4], befriending_seniors_list))
    befriending_seniors_list = [sublist[:2] + [sublist[3], sublist[2]] for sublist in befriending_seniors_list]  # Swap order of this week's update and previous week's
    table = add_table_with_grid(doc, rows=len(befriending_seniors_list)+1, cols=4)

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Senior'
    header_cells[0].paragraphs[0].runs[0].bold = True
    header_cells[1].text = 'Unit Number'
    header_cells[1].paragraphs[0].runs[0].bold = True
    header_cells[2].text = "Previous visit's Updates"
    header_cells[2].paragraphs[0].runs[0].bold = True
    header_cells[3].text = "This week's Visit Updates"
    header_cells[3].paragraphs[0].runs[0].bold = True

    for i in range(1,len(befriending_seniors_list)+1):
        table.cell(i, 0).text = befriending_seniors_list[i-1][0] # Senior Name
        table.cell(i, 1).text = befriending_seniors_list[i-1][1] # Unit Number
        table.cell(i, 2).text = befriending_seniors_list[i-1][2] # Last Week's Update
        table.cell(i, 3).text = befriending_seniors_list[i-1][3] # This Week's Update

    doc.add_page_break()
    title = doc.add_paragraph("Frail Seniors' House Visit Updates")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    frail_seniors_list = get_list_of_frail_seniors_status()[1:]
    frail_seniors_list = list(map(lambda row: row[:4], frail_seniors_list))
    frail_seniors_name_transformed = list(map(lambda item: extract_frail_senior_name_from_frail_list(item),
	 map(lambda row: row[0], frail_seniors_list)))
    seniors_status = list(map(lambda row: row[1:], frail_seniors_list))
    frail_seniors_list = [[name] + sublist for name, sublist in zip(frail_seniors_name_transformed, seniors_status)]
    frail_seniors_list = [sublist[:2] + [sublist[3], sublist[2]] for sublist in frail_seniors_list] # Swap order of this week's update and previous week's
    
    table = add_table_with_grid(doc, rows=len(frail_seniors_list)+1, cols=4)
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Senior'
    header_cells[0].paragraphs[0].runs[0].bold = True
    header_cells[1].text = 'Unit Number'
    header_cells[1].paragraphs[0].runs[0].bold = True
    header_cells[2].text = "Previous visit's Updates"
    header_cells[2].paragraphs[0].runs[0].bold = True
    header_cells[3].text = "This week's Visit Updates"
    header_cells[3].paragraphs[0].runs[0].bold = True

    for i in range(1,len(frail_seniors_list)+1):
        table.cell(i, 0).text = frail_seniors_list[i-1][0] # Senior Name
        table.cell(i, 1).text = frail_seniors_list[i-1][1] # Unit Number
        table.cell(i, 2).text = frail_seniors_list[i-1][2] # Last Week's Update
        table.cell(i, 3).text = frail_seniors_list[i-1][3] # This Week's Update

    date_of_visit = get_list_of_befriending_seniors_status()[0][2]
    doc.save(f"reports/{date_of_visit}_weekly_report.docx")